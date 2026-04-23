import streamlit as st
import anthropic
from io import BytesIO
import pypdf
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.set_page_config(
    page_title="Tailored Resume Builder",
    page_icon="📄",
    layout="wide",
)

# ── Styles ──────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 12px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .phase-badge {
        background: #667eea;
        color: white;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: bold;
    }
    .confidence-high { color: #28a745; font-weight: bold; }
    .confidence-mid  { color: #ffc107; font-weight: bold; }
    .confidence-low  { color: #dc3545; font-weight: bold; }
    .stTextArea textarea { font-family: monospace; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>📄 Tailored Resume Builder</h1>
    <p>AI-powered resume generation — research, discover, and tailor to any job description</p>
</div>
""", unsafe_allow_html=True)


# ── Helpers ──────────────────────────────────────────────────────────────────

def get_client(api_key: str) -> anthropic.Anthropic:
    return anthropic.Anthropic(api_key=api_key)


def _add_bottom_border(paragraph):
    """Add a bottom border line under a paragraph (used for section headers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_para_spacing(paragraph, before=0, after=0):
    pPr = paragraph._p.get_or_add_pPr()
    pSpacing = OxmlElement("w:spacing")
    pSpacing.set(qn("w:before"), str(before))
    pSpacing.set(qn("w:after"), str(after))
    pPr.append(pSpacing)


def _add_run_with_inline(para, text, base_size=Pt(10.5), base_font="Calibri"):
    """Render **bold**, *italic*, and plain text runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)
        run.font.name = base_font
        run.font.size = base_size


def markdown_to_docx(md_text: str, job_description: str = "") -> bytes:
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.85)
        sec.right_margin  = Inches(0.85)

    # Reset Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

    # ── Step 1: Strip unwanted sections from the docx output ─────────────────
    _skip_headings = re.compile(
        r"^#+\s*("
        r"(updated\s+)?match\s+analysis"
        r"|content\s+match"
        r"|selection\s+rationale"
        r"|certifications?"
        r"|gap\s+report"
        r"|unaddressed.*requirements?"
        r"|weakly\s+addressed"
        r"|overall\s+assessment"
        r")",
        re.IGNORECASE,
    )
    clean_lines = []
    skip = False
    for ln in md_text.splitlines():
        if _skip_headings.match(ln.strip()):
            skip = True
            continue
        if skip and re.match(r"^#+\s", ln.strip()):
            if not _skip_headings.match(ln.strip()):
                skip = False
        if not skip:
            clean_lines.append(ln)

    # ── Step 2: Filter & consolidate Skills section ───────────────────────────
    # Build a set of JD words for relevance matching
    jd_words = set(re.findall(r"[a-zA-Z][a-zA-Z0-9#+.\-]{1,}", job_description.lower())) if job_description else set()

    def _skill_relevant(skill_item: str) -> bool:
        """Return True if the skill token appears in the JD (case-insensitive)."""
        if not jd_words:
            return True  # no JD provided — keep everything
        token = skill_item.strip().lower()
        # Direct word match OR the token is a substring of any JD word
        return token in jd_words or any(token in jw or jw in token for jw in jd_words if len(jw) > 2)

    # Locate the SKILLS section and rewrite it with only 2 headings
    _skills_heading = re.compile(r"^##\s+skills", re.IGNORECASE)
    _next_h2 = re.compile(r"^##\s+", re.IGNORECASE)
    _skill_line = re.compile(r"^\*\*(.+?):\*\*\s*(.+)$")

    skills_start = None
    skills_end = None
    for idx, ln in enumerate(clean_lines):
        if _skills_heading.match(ln.strip()):
            skills_start = idx
        elif skills_start is not None and skills_end is None and _next_h2.match(ln.strip()) and idx > skills_start:
            skills_end = idx
            break
    if skills_start is not None and skills_end is None:
        skills_end = len(clean_lines)

    if skills_start is not None:
        # Collect all skill items from every sub-heading in the section
        all_skills = []
        for ln in clean_lines[skills_start + 1: skills_end]:
            m = _skill_line.match(ln.strip())
            if m:
                items = [s.strip() for s in m.group(2).split(",") if s.strip()]
                all_skills.extend(items)

        # Filter to JD-relevant only
        filtered = [s for s in all_skills if _skill_relevant(s)]
        if not filtered:
            filtered = all_skills  # fallback: keep all if filter removes everything

        # Split into 2 balanced groups
        mid = (len(filtered) + 1) // 2
        group1 = ", ".join(filtered[:mid])
        group2 = ", ".join(filtered[mid:])

        new_skills_block = ["## SKILLS"]
        new_skills_block.append(f"**Technical Skills:** {group1}")
        if group2:
            new_skills_block.append(f"**Tools & Platforms:** {group2}")

        clean_lines = clean_lines[:skills_start] + new_skills_block + clean_lines[skills_end:]

    lines = clean_lines
    _bullet_count = 0   # resets at each new H3 (job entry)

    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── H1: Candidate name (large, centred, dark) ────────────────────────
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=0, after=40)
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

        # ── H2: Section headers (blue with bottom border) ────────────────────
        elif stripped.startswith("## "):
            _bullet_count = 0
            text = stripped[3:].strip().upper()
            p = doc.add_paragraph()
            _set_para_spacing(p, before=120, after=20)
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            _add_bottom_border(p)

        # ── H3: Job title / sub-heading ──────────────────────────────────────
        elif stripped.startswith("### "):
            _bullet_count = 0   # new experience block — reset bullet counter
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            _set_para_spacing(p, before=80, after=0)
            _add_run_with_inline(p, text, base_size=Pt(10.5))
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)

        # ── Bullet points ────────────────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            # Rule: max 4 bullets per experience block
            if _bullet_count >= 4:
                i += 1
                continue
            _bullet_count += 1

            content = stripped[2:]
            # Remove em dashes and semicolons from bullet text
            content = content.replace("—", " ").replace("–", " ").replace(";", ",")
            content = re.sub(r" {2,}", " ", content).strip()

            p = doc.add_paragraph()
            _set_para_spacing(p, before=0, after=20)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "360")
            ind.set(qn("w:hanging"), "180")
            pPr.append(ind)
            bullet_run = p.add_run("• ")
            bullet_run.font.name = "Calibri"
            bullet_run.font.size = Pt(10.5)
            _add_run_with_inline(p, content)

        # ── Skill category line: **Label:** items ────────────────────────────
        elif re.match(r"^\*\*[^*]+:\*\*\s+.+", stripped):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=30, after=30)
            m = re.match(r"^\*\*([^*]+):\*\*\s+(.+)$", stripped)
            if m:
                label = p.add_run(m.group(1) + ": ")
                label.bold = True
                label.font.name = "Calibri"
                label.font.size = Pt(10.5)
                label.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
                items = p.add_run(m.group(2))
                items.font.name = "Calibri"
                items.font.size = Pt(10.5)
                items.font.color.rgb = RGBColor(0x26, 0x26, 0x26)
            else:
                _add_run_with_inline(p, stripped)

        # ── Horizontal rule → thin spacer ────────────────────────────────────
        elif re.match(r"^---+$", stripped) or re.match(r"^___+$", stripped):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=0, after=0)

        # ── Empty line → small vertical gap ──────────────────────────────────
        elif stripped == "":
            if i > 0 and lines[i - 1].strip() != "":
                p = doc.add_paragraph()
                _set_para_spacing(p, before=0, after=30)

        # ── Plain paragraph ───────────────────────────────────────────────────
        else:
            if stripped:
                _bullet_count = 0  # any new content block resets the per-entry bullet limit
            p = doc.add_paragraph()
            _set_para_spacing(p, before=0, after=16)
            _add_run_with_inline(p, stripped)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_system_prompt(resumes_text: str) -> str:
    has_resume = resumes_text.strip() != ""
    resume_section = f"""You have access to the user's resume library:
<resume_library>
{resumes_text}
</resume_library>""" if has_resume else """No resume was provided. Generate a complete, realistic resume from scratch based solely on the job description and keyword list. Invent plausible but realistic experience, roles, and achievements that match the JD perfectly."""

    return f"""You are an expert resume builder. Your goal is to produce the most JD-aligned resume possible.

{resume_section}

## Resume Generation Rules

You will receive a pre-extracted keyword list grouped into 6 buckets. Every keyword that fits the role MUST appear using the EXACT phrase from the list.

Placement rules per bucket:
- **Hard Skills** → Skills section + bullets
- **Tools** → Skills section + bullets
- **Business Terms** → bullets + summary
- **Functional Skills** → bullets + summary
- **Industry Terms** → summary + bullets
- **Preferred Keywords** → sprinkled across all sections

## Tool Substitution Rule (CRITICAL)
- If the candidate used a tool not mentioned in the JD, replace it with the closest equivalent tool from the JD keyword list
- Example: candidate used Python, JD only mentions Excel → write Excel, not Python
- This removes overqualification signals and aligns the resume to what the employer is looking for
- Apply this to every bullet where a non-JD tool appears

Do NOT rely on semantic similarity. If the keyword list says "data visualization", write "data visualization" exactly.

## Tool Substitution Rule
- If the candidate used a tool that does NOT appear in the keyword list, and the JD mentions a functionally equivalent tool for the same type of task, substitute the JD's tool in the bullet
- Keep the achievement, outcome, and numbers exactly the same — only the tool name changes
- Only substitute when the replacement tool can realistically perform the same task
- If no equivalent JD tool exists for that task, keep the candidate's original tool

## Tool Substitution Rule
- If the candidate used a tool that does NOT appear in the keyword list, and the JD mentions a functionally equivalent tool for the same type of task, substitute the JD's tool in the bullet
- Keep the achievement, outcome, and numbers exactly the same — only the tool name changes
- Only substitute when the replacement tool can realistically perform the same task
- If no equivalent JD tool exists for that task, keep the candidate's original tool

## Bullet Point Rules (STRICT — zero tolerance)
- Maximum 4 bullets per job/experience
- Maximum 28 words per bullet — generate short and punchy from the start, never exceed this
- Every bullet starts with a strong action verb from the JD's vocabulary
- ZERO em dashes (—) or en dashes (–) in bullets — rephrase instead
- ZERO semicolons (;) in bullets — use a comma or split into a new clause

## Page & Projects Rules
- Target ONE page — achieve this through the 28-word bullet limit and 4-bullet max, never by dropping bullets from any experience
- Every experience MUST have its bullets — do not leave any role without bullet points
- For the Projects section: use ONLY the top 3 projects ranked in the keyword extraction step — ignore all others
- Show the relevance score next to each project heading (e.g. "Project Name — 87/100")

Always respond in structured Markdown."""


# ── CALL 1: Extract & bucket JD keywords ────────────────────────────────────
def stream_extract_keywords(client, job_description: str, resumes_text: str = ""):
    system = """You are a precise keyword extraction specialist for ATS-optimised resumes.

Extract every meaningful keyword from the job description and group into exactly these 6 buckets:

1. **Hard Skills** — technical competencies (e.g. statistical analysis, data modeling, A/B testing)
2. **Tools** — specific software, platforms, languages, frameworks (e.g. Python, Tableau, SQL, Salesforce)
3. **Business Terms** — business/domain language and role-specific phrases (e.g. "go-to-market", "data integrity", "revenue impact")
4. **Functional Skills** — cross-functional competencies (e.g. stakeholder management, cross-functional collaboration, project delivery)
5. **Industry Terms** — industry-specific vocabulary and concepts (e.g. "self-service dashboards", "churn analysis", "ETL pipelines")
6. **Preferred Keywords** — repeated or emphasised phrases that should appear across the resume (e.g. "data-driven", "actionable insights")

Rules:
- Capture EXACT phrases as written in the JD — no paraphrasing
- Flag keywords that appear more than once with ⭐ (highest priority)
- Include every tool mentioned even if only once
- Do not skip Business Terms or Industry Terms — these are often missed and matter for ATS

Output format (Markdown, exactly as shown):
## Extracted Keywords

**Hard Skills:** keyword1, keyword2 ⭐, keyword3
**Tools:** tool1 ⭐, tool2, tool3
**Business Terms:** term1, term2 ⭐
**Functional Skills:** skill1 ⭐, skill2
**Industry Terms:** term1, term2
**Preferred Keywords:** phrase1 ⭐, phrase2

## Project Relevance Ranking

Rank ALL projects found in the candidate's resume by relevance to this JD. Score each out of 100.

| Rank | Project Name | Relevance Score | Why |
|------|-------------|-----------------|-----|
| 1 | Project A | 92/100 | Directly uses X and Y from JD |
| 2 | Project B | 74/100 | Covers Z competency |
| 3 | Project C | 61/100 | Partial match on ... |

Only the top 3 will be included in the resume."""

    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=1200,
        system=system,
        messages=[{"role": "user", "content": f"Job Description:\n{job_description}\n\nCandidate Resume Library (for project ranking):\n{resumes_text}"}],
    ) as stream:
        for text in stream.text_stream:
            yield text


# ── CALL 2: Generate resume anchored to keyword list ────────────────────────
def stream_tailored_resume(client, system_prompt: str, job_description: str,
                            keyword_list: str, extra_context: str):
    user_message = f"""Build a tailored resume using the candidate's experience library, the pre-extracted keyword list, and the project relevance rankings below.

## Pre-Extracted Keyword List + Project Rankings (MANDATORY)
{keyword_list}

## Job Description
{job_description}
"""
    if extra_context.strip():
        user_message += f"\n## Additional Candidate Context\n{extra_context}\n"

    user_message += """
Rules:
- Use ONLY the top 3 ranked projects from the Project Relevance Ranking table — ignore all others
- Show the relevance score next to each project heading e.g. "Project Name — 87/100"
- Keep every bullet under 28 words — punchy and tight from the start
- Every experience must have bullets — never leave a role empty
- Fit one page through brevity, not by dropping bullets

Produce:
1. The full tailored resume in Markdown
2. A gap report listing any unaddressed JD requirements with mitigation tips"""

    full_response = ""
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            full_response += text
            yield text

    return full_response


# ── CALL 3: Exact phrase audit + patch ──────────────────────────────────────
def stream_keyword_audit(client, resume_md: str, keyword_list: str, job_description: str):
    system = """You are a strict ATS keyword auditor. Your job is to check a resume against a keyword list and patch any missing exact phrases.

Rules:
- Compare the resume word-for-word against every keyword in the list
- If an exact keyword phrase is missing, insert it naturally and meaningfully
- The sentence must still read as a clear, specific achievement after insertion — not a keyword dump
- Insert into: the closest matching bullet, the summary, or the skills line — whichever fits best
- If a keyword cannot be inserted without making the sentence awkward or meaningless, skip it
- Tool substitution: if a bullet mentions a tool NOT in the keyword list, replace it with the closest equivalent tool that IS in the keyword list
- Do NOT change any names, dates, companies, or job titles
- Do NOT add new bullets — only edit existing text
- Preserve all bullet rules: max 4 bullets per role, no em dashes, no semicolons
- ZERO TOLERANCE: Every bullet must be 28 words or fewer — count the words, rewrite until it fits
- Do NOT truncate — rewrite intelligently to keep the keyword, metric, and core impact
- This is mandatory: no bullet may leave this step longer than 28 words under any circumstance
- Output ONLY the patched resume in Markdown — no commentary"""

    user_message = f"""Audit this resume against the keyword list. Insert any missing exact phrases naturally.

## Keyword List
{keyword_list}

## Resume to Audit
{resume_md}

## Job Description (for context)
{job_description}"""

    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=system,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            yield text


def stream_experience_discovery(client, system_prompt: str, job_description: str, gap_context: str) -> str:
    user_message = f"""Based on the job description and the gap report below, ask me targeted branching questions to surface undocumented experiences that could fill these gaps.

**Job Description:**
{job_description}

**Gap Report / Context:**
{gap_context}

Ask 3-5 specific, concrete questions. After I answer, you will incorporate my answers into a revised resume."""

    full_response = ""
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            full_response += text
            yield text


def stream_ai_suggested_answers(client, job_description: str, questions: str, resume_context: str):
    system = """You are playing the role of the single most ideal candidate this company could hire.

Before answering, you deeply analyze the job description to extract:
- EXACT tools, software, and technologies mentioned (e.g. SQL, Tableau, dbt, Salesforce, Python, Jira)
- EXACT methodologies and frameworks mentioned (e.g. Agile, A/B testing, OKRs, ETL pipelines)
- EXACT language and phrasing the company uses to describe the work
- EXACT outcomes and metrics the company cares about (revenue, retention, accuracy, speed, scale)
- The company's industry domain and the specific way they do things in that domain

Then you answer every discovery question as someone who:
- Used those EXACT tools in their answers (not similar ones — the actual ones from the JD)
- Followed those EXACT methodologies the company values
- Achieved outcomes using the EXACT metrics the company tracks
- Speaks using the EXACT terminology and vocabulary the company uses in the JD
- Has done the work in the EXACT way the company describes it — not a generic version

Rules:
- Never use vague tool names. If the JD says "dbt" say dbt, not "a data transformation tool"
- Never use generic verbs. If the JD says "orchestrate" use orchestrate, not "manage"
- Every answer must reference at least one specific tool or technology from the JD
- Numbers and scale must match what the company would find impressive for this role level
- Answer in first person, naturally, like a real human recalling a real experience

Format: Answer each question numbered to match. Be specific, concrete, and role-precise."""

    user_message = f"""Study this job description carefully — extract every tool, technology, methodology, metric, and piece of company vocabulary before answering:

<job_description>
{job_description}
</job_description>

Candidate's existing resume for context (build on their background, don't contradict it):
<resume>
{resume_context}
</resume>

Now answer each discovery question AS the ideal candidate for this exact role — using the exact tools, exact methods, exact language, and exact outcomes this company is looking for:

{questions}

Remember: use their tools, their words, their metrics. Make every answer feel like it was written by someone who has lived inside this company's world."""

    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=2000,
        system=system,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            yield text


def stream_humanized_resume(client, resume_md: str, job_description: str):
    system = """You are a professional resume editor with two goals: make the resume sound natural and human, AND make sure every bullet actively uses the language, keywords, and terminology from the job description.

Before rewriting, scan the job description and extract:
- Every specific tool, technology, and software mentioned
- Every action verb and methodology the JD uses
- Every outcome, metric, or result the role cares about
- The exact phrasing and vocabulary the company uses

Then rewrite each bullet to:
1. WEAVE IN the JD's exact keywords, tools, and language naturally — do not just preserve what is already there, actively work them in where truthful and relevant
2. MIRROR the JD's terminology — if the JD says "cross-functional collaboration" use that phrase, not "worked with teams"
3. Sound natural and confident — no stiff corporate jargon
4. REMOVE all em dashes (—) and en dashes (–) — rephrase instead
5. REMOVE all semicolons (;) — use commas or split into a new clause
6. Start every bullet with a strong action verb from the JD's vocabulary where possible
7. Keep the same Markdown structure (headings, bullets, sections) — do not add or remove sections
8. Do not change names, dates, companies, job titles, or numbers
9. Output ONLY the rewritten resume in Markdown — no commentary, no analysis"""

    user_message = f"""Rewrite this resume to sound human AND to actively use the job description's exact keywords and language in every bullet.

Job Description — extract all keywords, tools, verbs, and terminology from this first:
<job_description>
{job_description}
</job_description>

Resume to rewrite:
<resume>
{resume_md}
</resume>

Remember: do not just preserve existing keywords — actively weave JD language into bullets where truthful. Remove all em dashes and semicolons."""

    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=system,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            yield text


def stream_gap_report(client, resume_md: str, job_description: str):
    system = """You are a resume gap analyst. Given a resume and a job description, identify requirements from the JD that are unaddressed or weakly addressed in the resume. Be concise and actionable."""

    user_message = f"""Analyze this resume against the job description and produce a Gap Report.

Job Description:
{job_description}

Resume:
{resume_md}

Output a Gap Report covering:
1. Unaddressed or weakly addressed requirements (with severity: High / Medium / Low)
2. Mitigation strategy for each gap
3. Overall assessment (1–2 sentences)"""

    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=1500,
        system=system,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            yield text


def stream_revised_resume(client, system_prompt: str, job_description: str, original_resume: str, discovery_qa: str) -> str:
    user_message = f"""Revise the tailored resume below by incorporating the newly discovered experiences from our Q&A session.

**Job Description:**
{job_description}

**Previously Generated Resume:**
{original_resume}

**Newly Discovered Experiences (Q&A):**
{discovery_qa}

Produce the final revised resume and an updated match analysis."""

    full_response = ""
    with client.messages.stream(
        model="claude-sonnet-4-6",
        max_tokens=4000,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        for text in stream.text_stream:
            full_response += text
            yield text


# ── Session state init ────────────────────────────────────────────────────────
for key, default in {
    "phase": "input",
    "tailored_resume": "",
    "gap_context": "",
    "discovery_questions": "",
    "discovery_answers": "",
    "ai_suggested_answers": "",
    "final_resume": "",
    "humanized_resume": "",
    "humanized_gap": "",
    "system_prompt": "",
    "job_description": "",
    "pre_humanize_phase": "done",
    "keyword_list": "",          # extracted keyword buckets from Call 1
    "draft_resume": "",          # raw output from Call 2 before audit
}.items():
    if key not in st.session_state:
        st.session_state[key] = default


# ── Resolve API key (secrets → sidebar fallback) ─────────────────────────────
api_key = st.secrets.get("ANTHROPIC_API_KEY", "") if hasattr(st, "secrets") else ""

# ── Sidebar — resume upload (+ key fallback if no secret set) ────────────────
with st.sidebar:
    if not api_key:
        st.header("⚙️ Configuration")
        api_key = st.text_input("Anthropic API Key", type="password", placeholder="sk-ant-...")
        st.caption("Your key is never stored — it lives only in this session.")
        st.divider()

    st.header("📂 Resume Library")
    uploaded_files = st.file_uploader(
        "Upload your existing resumes (optional)",
        accept_multiple_files=True,
        type=["txt", "md", "pdf"],
        help="Optional — if no resume is uploaded, a full resume is generated from scratch using the JD.",
    )

    def extract_text(file) -> str:
        if file.name.lower().endswith(".pdf"):
            reader = pypdf.PdfReader(BytesIO(file.read()))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        return file.read().decode("utf-8", errors="replace")

    resumes_text = ""
    if uploaded_files:
        st.success(f"{len(uploaded_files)} resume(s) loaded")
        for i, f in enumerate(uploaded_files, 1):
            content = extract_text(f)
            resumes_text += f"\n\n--- RESUME {i}: {f.name} ---\n{content}"

    st.divider()
    if st.button("🔄 Start Over", use_container_width=True):
        for key in ["phase", "tailored_resume", "gap_context", "discovery_questions",
                    "discovery_answers", "ai_suggested_answers", "final_resume", "system_prompt",
                    "job_description", "keyword_list", "draft_resume"]:
            st.session_state[key] = "" if key != "phase" else "input"
        st.rerun()


# ── Main area ─────────────────────────────────────────────────────────────────

# ── PHASE: INPUT ──────────────────────────────────────────────────────────────
if st.session_state.phase == "input":
    col1, col2 = st.columns([3, 2])

    with col1:
        st.subheader("📋 Job Description")
        job_description = st.text_area(
            "Paste the full job description here",
            height=350,
            placeholder="Paste the job description (including requirements, responsibilities, company info)...",
        )

    with col2:
        st.subheader("💬 Additional Context (optional)")
        extra_context = st.text_area(
            "Any experiences, projects, or background not in your resumes?",
            height=180,
            placeholder="E.g. I led a Kubernetes migration project for a non-profit last year but it's not in my resumes...",
        )
        st.info("**Tips for best results:**\n- Upload at least 1 resume\n- Include a detailed JD\n- Add any recent experiences not yet in your resumes")

    st.divider()
    ready = api_key and job_description.strip()
    if not resumes_text:
        st.info("No resume uploaded — a complete resume will be generated from scratch based on the JD.")
    if not ready:
        missing = []
        if not job_description.strip(): missing.append("a job description")
        if not api_key:       missing.append("an API Key (sidebar)")
        st.warning(f"Please provide: {', '.join(missing)}")

    if st.button("🚀 Tailor My Resume", disabled=not ready, use_container_width=True, type="primary"):
        st.session_state.system_prompt = build_system_prompt(resumes_text)
        st.session_state.job_description = job_description
        st.session_state.extra_context = extra_context
        st.session_state.phase = "tailoring"
        st.rerun()


# ── PHASE: TAILORING ──────────────────────────────────────────────────────────
elif st.session_state.phase == "tailoring":
    st.subheader("⚡ Building Your Tailored Resume")

    client = get_client(api_key)

    try:
        # ── Call 1: Extract keywords ──────────────────────────────────────────
        if not st.session_state.keyword_list:
            st.markdown("**Step 1 of 3 — Extracting JD keywords into buckets...**")
            placeholder1 = st.empty()
            kw_text = ""
            with st.spinner("Analysing job description..."):
                for chunk in stream_extract_keywords(client, st.session_state.job_description, st.session_state.system_prompt):
                    kw_text += chunk
                    placeholder1.markdown(kw_text)
            st.session_state.keyword_list = kw_text
            st.rerun()

        # ── Call 2: Generate resume anchored to keywords ──────────────────────
        if st.session_state.keyword_list and not st.session_state.draft_resume:
            st.markdown("**Step 2 of 3 — Building resume from keyword-anchored experience matching...**")
            placeholder2 = st.empty()
            draft_text = ""
            with st.spinner("Generating resume..."):
                for chunk in stream_tailored_resume(
                    client,
                    st.session_state.system_prompt,
                    st.session_state.job_description,
                    st.session_state.keyword_list,
                    st.session_state.get("extra_context", ""),
                ):
                    draft_text += chunk
                    placeholder2.markdown(draft_text)
            st.session_state.draft_resume = draft_text
            st.rerun()

        # ── Call 3: Keyword audit + patch ─────────────────────────────────────
        if st.session_state.draft_resume and not st.session_state.tailored_resume:
            st.markdown("**Step 3 of 3 — Running exact keyword audit and patching gaps...**")
            placeholder3 = st.empty()
            final_text = ""
            with st.spinner("Auditing keyword coverage..."):
                for chunk in stream_keyword_audit(
                    client,
                    st.session_state.draft_resume,
                    st.session_state.keyword_list,
                    st.session_state.job_description,
                ):
                    final_text += chunk
                    placeholder3.markdown(final_text)
            st.session_state.tailored_resume = final_text
            st.session_state.phase = "review"
            st.rerun()

    except anthropic.AuthenticationError:
        st.error("Invalid API key. Please check your Anthropic API key in the sidebar.")
        st.session_state.phase = "input"
    except Exception as e:
        st.error(f"Error: {e}")
        st.session_state.phase = "input"


# ── PHASE: REVIEW ─────────────────────────────────────────────────────────────
elif st.session_state.phase == "review":
    st.subheader("✅ Phase 1 Complete — Review Your Tailored Resume")

    tab1, tab2 = st.tabs(["📄 Resume Preview", "📝 Raw Markdown"])
    with tab1:
        st.markdown(st.session_state.tailored_resume)
    with tab2:
        st.code(st.session_state.tailored_resume, language="markdown")

    st.divider()
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.download_button(
            "⬇️ Download (.docx)",
            data=markdown_to_docx(st.session_state.tailored_resume, st.session_state.job_description),
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    with col2:
        if st.button("🔍 Experience Discovery", use_container_width=True):
            st.session_state.gap_context = st.session_state.tailored_resume
            st.session_state.phase = "discovery"
            st.rerun()

    with col3:
        if st.button("🪄 Humanize Resume", use_container_width=True):
            st.session_state.pre_humanize_phase = "review"
            st.session_state.humanized_resume = ""
            st.session_state.humanized_gap = ""
            st.session_state.phase = "humanizing"
            st.rerun()

    with col4:
        if st.button("✅ Done!", use_container_width=True, type="primary"):
            st.session_state.final_resume = st.session_state.tailored_resume
            st.session_state.phase = "done"
            st.rerun()


# ── PHASE: DISCOVERY ──────────────────────────────────────────────────────────
elif st.session_state.phase == "discovery":
    st.subheader("🔍 Phase 2 — Experience Discovery")
    st.caption("Claude will ask targeted questions to surface undocumented experiences that fill gaps in your resume.")

    if not st.session_state.discovery_questions:
        client = get_client(api_key)
        output_placeholder = st.empty()
        full_text = ""
        with st.spinner("Generating discovery questions..."):
            for chunk in stream_experience_discovery(
                client,
                st.session_state.system_prompt,
                st.session_state.job_description,
                st.session_state.gap_context,
            ):
                full_text += chunk
                output_placeholder.markdown(full_text)
        st.session_state.discovery_questions = full_text
        st.rerun()
    else:
        st.markdown(st.session_state.discovery_questions)

    st.divider()

    # ── AI Answer Suggester ───────────────────────────────────────────────────
    st.markdown("### ✍️ Your Answers")
    st.caption("Write your own answers, or let AI suggest ideal answers based on the job description — then edit as needed.")

    col_btn1, col_btn2 = st.columns([1, 3])
    with col_btn1:
        suggest_clicked = st.button("🤖 Suggest AI Answers", use_container_width=True)

    if suggest_clicked:
        client = get_client(api_key)
        suggest_placeholder = st.empty()
        suggested = ""
        with st.spinner("AI is generating ideal answers for this role..."):
            for chunk in stream_ai_suggested_answers(
                client,
                st.session_state.job_description,
                st.session_state.discovery_questions,
                st.session_state.gap_context,
            ):
                suggested += chunk
                suggest_placeholder.markdown(suggested)
        st.session_state.ai_suggested_answers = suggested
        st.rerun()

    if st.session_state.ai_suggested_answers:
        st.info("**AI-suggested answers below** — these are crafted to match the job description. Edit, remove, or replace anything that doesn't apply to you.")

    answers = st.text_area(
        "Your answers (edit AI suggestions or write your own):",
        height=350,
        placeholder="Answer each question as specifically as possible...\n\nTip: Click '🤖 Suggest AI Answers' above to get a starting point.",
        value=st.session_state.ai_suggested_answers or st.session_state.discovery_answers,
    )

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ Skip — keep current resume", use_container_width=True):
            st.session_state.phase = "review"
            st.rerun()
    with col2:
        if st.button("🔄 Revise Resume with My Answers", disabled=not answers.strip(),
                     use_container_width=True, type="primary"):
            st.session_state.discovery_answers = answers
            st.session_state.phase = "revision"
            st.rerun()


# ── PHASE: REVISION ───────────────────────────────────────────────────────────
elif st.session_state.phase == "revision":
    st.subheader("🔄 Phase 3 — Revising Resume with Discovered Experiences")

    client = get_client(api_key)
    output_placeholder = st.empty()
    full_text = ""

    discovery_qa = (
        f"Questions:\n{st.session_state.discovery_questions}\n\n"
        f"Answers:\n{st.session_state.discovery_answers}"
    )

    try:
        with st.spinner("Revising resume..."):
            for chunk in stream_revised_resume(
                client,
                st.session_state.system_prompt,
                st.session_state.job_description,
                st.session_state.tailored_resume,
                discovery_qa,
            ):
                full_text += chunk
                output_placeholder.markdown(full_text)

        st.session_state.final_resume = full_text
        st.session_state.phase = "done"
        st.rerun()

    except Exception as e:
        st.error(f"Error: {e}")
        st.session_state.phase = "review"


# ── PHASE: DONE ───────────────────────────────────────────────────────────────
elif st.session_state.phase == "done":
    st.success("🎉 Your tailored resume is ready!")

    tab1, tab2 = st.tabs(["📄 Final Resume", "📝 Raw Markdown"])
    with tab1:
        st.markdown(st.session_state.final_resume)
    with tab2:
        st.code(st.session_state.final_resume, language="markdown")

    st.divider()
    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button(
            "⬇️ Download Final Resume (.docx)",
            data=markdown_to_docx(st.session_state.final_resume, st.session_state.job_description),
            file_name="tailored_resume_final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
    with col2:
        if st.button("🪄 Humanize Resume", use_container_width=True):
            st.session_state.pre_humanize_phase = "done"
            st.session_state.humanized_resume = ""
            st.session_state.humanized_gap = ""
            st.session_state.phase = "humanizing"
            st.rerun()
    with col3:
        if st.button("🔁 Tailor for another job", use_container_width=True):
            for key in ["phase", "tailored_resume", "gap_context", "discovery_questions",
                        "discovery_answers", "final_resume", "humanized_resume", "humanized_gap"]:
                st.session_state[key] = "" if key != "phase" else "input"
            st.session_state.phase = "input"
            st.rerun()


# ── PHASE: HUMANIZING ─────────────────────────────────────────────────────────
elif st.session_state.phase == "humanizing":
    st.subheader("🪄 Humanizing Your Resume")

    # Decide which resume to humanize
    source = (st.session_state.final_resume
              if st.session_state.pre_humanize_phase == "done"
              else st.session_state.tailored_resume)

    client = get_client(api_key)

    # ── Step 1: Humanize ──────────────────────────────────────────────────────
    if not st.session_state.humanized_resume:
        st.markdown("**Step 1 of 2 — Rewriting for natural language...**")
        placeholder = st.empty()
        full_text = ""
        with st.spinner("Humanizing resume — preserving all keywords..."):
            for chunk in stream_humanized_resume(client, source, st.session_state.job_description):
                full_text += chunk
                placeholder.markdown(full_text)
        st.session_state.humanized_resume = full_text
        st.rerun()

    # ── Step 2: Re-run gap report ─────────────────────────────────────────────
    if st.session_state.humanized_resume and not st.session_state.humanized_gap:
        st.markdown("**Step 2 of 2 — Re-running gap analysis...**")
        placeholder2 = st.empty()
        gap_text = ""
        with st.spinner("Running gap analysis on humanized resume..."):
            for chunk in stream_gap_report(client, st.session_state.humanized_resume,
                                           st.session_state.job_description):
                gap_text += chunk
                placeholder2.markdown(gap_text)
        st.session_state.humanized_gap = gap_text
        st.session_state.phase = "humanized"
        st.rerun()


# ── PHASE: HUMANIZED ──────────────────────────────────────────────────────────
elif st.session_state.phase == "humanized":
    st.success("🪄 Humanized resume ready!")

    tab1, tab2, tab3 = st.tabs(["📄 Humanized Resume", "📊 Gap Report", "📝 Raw Markdown"])
    with tab1:
        st.markdown(st.session_state.humanized_resume)
    with tab2:
        st.markdown(st.session_state.humanized_gap)
    with tab3:
        st.code(st.session_state.humanized_resume, language="markdown")

    st.divider()
    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button(
            "⬇️ Download Humanized Resume (.docx)",
            data=markdown_to_docx(st.session_state.humanized_resume, st.session_state.job_description),
            file_name="humanized_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )

    with col2:
        back_phase = st.session_state.pre_humanize_phase
        if st.button("⬅️ Back to Resume", use_container_width=True):
            st.session_state.phase = back_phase
            st.rerun()

    with col3:
        if st.button("🔁 Tailor for another job", use_container_width=True):
            for key in ["phase", "tailored_resume", "gap_context", "discovery_questions",
                        "discovery_answers", "final_resume", "humanized_resume", "humanized_gap"]:
                st.session_state[key] = "" if key != "phase" else "input"
            st.session_state.phase = "input"
            st.rerun()
